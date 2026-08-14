# -*- coding: utf-8 -*-
"""
Markdown -> Word(.docx) 转换器
支持: 标题/表格/代码块/无序有序列表/加粗/行内代码, 中文字体(宋体/黑体)排版
用法: py -3 md2docx.py <输入.md> [输出.docx]
"""
import sys
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name_latin="Calibri", name_ea="宋体", size=None, bold=None, color=None):
    run.font.name = name_latin
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name_ea)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def shade_paragraph(p, fill="F2F2F2"):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def add_inline(par, text, base_bold=False):
    """解析 **bold** 与 `code`"""
    tokens = re.split(r'(\*\*.+?\*\*|`[^`]+`)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = par.add_run(tok[2:-2])
            set_font(r, size=10.5, bold=True)
        elif tok.startswith('`') and tok.endswith('`'):
            r = par.add_run(tok[1:-1])
            set_font(r, name_latin="Consolas", name_ea="宋体", size=9.5)
        else:
            r = par.add_run(tok)
            set_font(r, size=10.5, bold=base_bold)


def add_code_block(doc, code):
    for line in code.rstrip('\n').split('\n'):
        p = doc.add_paragraph()
        r = p.add_run(line if line else ' ')
        set_font(r, name_latin="Consolas", name_ea="宋体", size=9)
        shade_paragraph(p)


def add_table(doc, rows):
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.paragraphs[0].text = ''
            par = cell.paragraphs[0]
            add_inline(par, cell_text.strip(), base_bold=(i == 0))
            if i == 0:
                shade_paragraph(par, fill="DEEAF6")
    doc.add_paragraph()


def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    # 全局样式: 中文字体
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    i = 0
    in_code = False
    code_buf = []
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()

        # 代码块
        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, '\n'.join(code_buf))
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 表格
        if line.startswith('|') and i + 1 < n and re.match(r'^\s*\|[\s:|-]+\|\s*$', lines[i + 1]):
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                rows.append(cells)
                i += 1
            add_table(doc, rows)
            continue

        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            h = doc.add_heading(level=min(level, 4))
            h.clear()
            r = h.add_run(text)
            set_font(r, name_latin="Calibri", name_ea="黑体", bold=True,
                     size={1: 16, 2: 14, 3: 12, 4: 11}.get(level, 11),
                     color=(0x1F, 0x3B, 0x63))
            i += 1
            continue

        # 分隔线
        if re.match(r'^\s*---+\s*$', line):
            i += 1
            continue

        # 有序/无序列表
        m = re.match(r'^\s*[-*]\s+(.*)$', line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, m.group(1))
            i += 1
            continue
        m = re.match(r'^\s*\d+\.\s+(.*)$', line)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, m.group(1))
            i += 1
            continue

        # 普通段落
        if line.strip():
            p = doc.add_paragraph()
            add_inline(p, line)
        i += 1

    doc.save(docx_path)
    print('已生成:', docx_path)


if __name__ == '__main__':
    md = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else md.rsplit('.', 1)[0] + '.docx'
    convert(md, out)
